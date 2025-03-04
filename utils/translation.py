# utils/translation.py

import os
import pysrt
from deep_translator import GoogleTranslator
from docx import Document

def translate_srt_file(srt_path, target_language='ne', suffix=None):
    """
    Translates the SRT file to the specified target language.

    Args:
        srt_path (str): Path to the input SRT file.
        target_language (str): Target language code (default is 'ne').
        suffix (str): Suffix to append to the output filename. If not provided,
                      defaults to f"_{target_language}".
    
    Returns:
        str: Path to the translated SRT file.
    """
    if suffix is None:
        suffix = f"_{target_language}"
        
    subs = pysrt.open(srt_path, encoding='utf-8')
    translator = GoogleTranslator(source='auto', target=target_language)
    
    for sub in subs:
        try:
            translated_text = translator.translate(sub.text)
            sub.text = translated_text
        except Exception as e:
            print(f"Error translating subtitle '{sub.text}': {e}")

    base, ext = os.path.splitext(srt_path)
    new_srt_path = f"{base}{suffix}{ext}"
    subs.save(new_srt_path, encoding='utf-8')
    print(f"Translated SRT saved to '{new_srt_path}'")
    return new_srt_path

def translate_paragraph(paragraph, translator):
    """
    Translates all text runs in a paragraph using the provided translator.
    """
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        if run.text and run.text.strip():
            try:
                translated_text = translator.translate(run.text)
                run.text = translated_text
            except Exception as e:
                print(f"Error translating run text '{run.text}': {e}")

def translate_document(file_path, target_language='ne'):
    """
    Translates a Word document to the specified target language.
    
    Args:
        file_path (str): Path to the input .docx file.
        target_language (str): Target language code (default is 'ne').

    Returns:
        Document: The translated Document object.
    """
    print(f"Translating document: {file_path}")
    document = Document(file_path)
    translator = GoogleTranslator(source='auto', target=target_language)
    
    for paragraph in document.paragraphs:
        translate_paragraph(paragraph, translator)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    translate_paragraph(paragraph, translator)
    return document
